#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
import re
import logging
from docx import Document
import unicodedata
import yaml
from pathlib import Path

# 配置日志 (可以根据 Flask 应用的日志配置调整或移除)
# logging.basicConfig( # Flask app 会处理日志配置
# level=logging.INFO,
# format='%(asctime)s - %(levelname)s - %(message)s',
# handlers=[logging.StreamHandler()]
# )
# logger = logging.getLogger(__name__)
# 使用 Flask 的 logger
logger = logging.getLogger('flask.app')


# 默认字数要求配置 (与之前脚本一致)
DEFAULT_REQUIREMENTS = {
    "摘要": {"max": 300},
    "权利要求书": {"min": 1500, "max": 2000},
    "说明书摘要": {"max": 300},
    "说明书": {
        "sub_sections": ["技术领域", "背景技术", "发明内容", "具体实施方式", "有益效果", "附图说明"],
        "min": 6000,
        "max": 10000
    },
    "技术领域": {"min": 50, "max": 300},
    "背景技术": {"min": 300, "max": 1000},
    "发明内容": {"min": 500, "max": 1500},
    "具体实施方式": {"ratio": 2.0, "reference": "权利要求书", "tolerance": 0.3, "min": 3000},
    "有益效果": {"min": 300, "max": 800},
    "附图说明": {"min": 50, "max": 500},
    "总字数": {"min": 9000, "max": 12000},
}

# 常见章节名称及其可能的正则表达式模式 (与之前脚本一致)
COMMON_SECTIONS_PATTERNS = {
    "权利要求书": [r"^\s*权\s*利\s*要\s*求\s*书\s*$", r"^\s*权\s*利\s*要\s*求\s*$"],
    "说明书摘要": [r"^\s*说\s*明\s*书\s*摘\s*要\s*$", r"^\s*摘\s*要\s*$"],
    "技术领域": [r"^\s*技\s*术\s*领\s*域\s*$", r"^\s*一、\s*技\s*术\s*领\s*域\s*$", r"^\s*1\.\s*技\s*术\s*领\s*域\s*"],
    "背景技术": [r"^\s*背\s*景\s*技\s*术\s*$", r"^\s*二、\s*背\s*景\s*技\s*术\s*$", r"^\s*2\.\s*背\s*景\s*技\s*术\s*"],
    "发明内容": [r"^\s*发\s*明\s*内\s*容\s*$", r"^\s*三、\s*发\s*明\s*内\s*容\s*$", r"^\s*3\.\s*发\s*明\s*内\s*容\s*"],
    "具体实施方式": [r"^\s*具\s*体\s*实\s*施\s*方\s*式\s*$", r"^\s*四、\s*具\s*体\s*实\s*施\s*方\s*式\s*$", r"^\s*4\.\s*具\s*体\s*实\s*施\s*方\s*式\s*"],
    "有益效果": [r"^\s*有\s*益\s*效\s*果\s*$", r"^\s*五、\s*有\s*益\s*效\s*果\s*$", r"^\s*5\.\s*有\s*益\s*效\s*果\s*"],
    "附图说明": [r"^\s*附\s*图\s*说\s*明\s*$", r"^\s*六、\s*附\s*图\s*说\s*明\s*$", r"^\s*6\.\s*附\s*图\s*说\s*明\s*"],
    "说明书": [r"^\s*说\s*明\s*书\s*$"]
}

class PatentAnalyzer:
    def __init__(self, file_path, config_data=None, count_mode="chinese"): # config_data is now a dict
        self.file_path = Path(file_path)
        self.count_mode = count_mode
        self.config = DEFAULT_REQUIREMENTS.copy() # Start with defaults
        self.paragraphs = []
        self.full_text_content = ""

        logger.info(f"PatentAnalyzer init: file_path='{self.file_path}', count_mode='{self.count_mode}'")

        if not self.file_path.exists():
            logger.error(f"文件不存在: {self.file_path}")
            raise FileNotFoundError(f"文件不存在: {self.file_path}")

        if config_data and isinstance(config_data, dict):
            # Merge user config with default, user config takes precedence for shared keys
            for key, value in config_data.items():
                if key in self.config and isinstance(self.config[key], dict) and isinstance(value, dict):
                    self.config[key].update(value)
                else:
                    self.config[key] = value
            logger.info(f"已加载并合并用户提供的配置数据。")
        else:
            logger.info("未提供有效配置数据或非字典格式，将使用默认字数要求配置。")


        self._load_content()

    def _load_content(self):
        logger.info(f"开始加载内容，文件路径: {self.file_path}")
        logger.info(f"文件是否存在: {self.file_path.exists()}")
        logger.info(f"文件名: {self.file_path.name}")
        logger.info(f"文件后缀 (来自 Path.suffix): '{self.file_path.suffix}'")

        ext = self.file_path.suffix.lower()
        logger.info(f"小写后缀用于判断: '{ext}'")

        try:
            if ext == ".docx":
                doc = Document(self.file_path)
                self.paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            elif ext == ".txt":
                with open(self.file_path, 'r', encoding='utf-8') as f:
                    self.paragraphs = [line.strip() for line in f if line.strip()]
            else:
                logger.error(f"不支持的文件类型检测到: '{ext}' (原始文件名: {self.file_path.name})")
                raise ValueError(f"不支持的文件类型: '{ext}'. 请提供 .txt 或 .docx 文件。")
            self.full_text_content = "\n".join(self.paragraphs)
            logger.info(f"已加载文档: {self.file_path} (共 {len(self.paragraphs)} 段)")
        except Exception as e:
            logger.error(f"加载文档 '{self.file_path}' 时失败: {e}", exc_info=True)
            raise

    def count_chars(self, text):
        if not text:
            return 0
        mode = self.count_mode
        if mode == "chinese":
            count = 0
            for char in text:
                if ('\u4e00' <= char <= '\u9fff' or
                    '\u3400' <= char <= '\u4dbf' or
                    '\U00020000' <= char <= '\U0002a6df'):
                    count += 1
            return count
        elif mode == "word":
            chinese_count = sum(1 for char in text if '\u4e00' <= char <= '\u9fff' or
                               '\u3400' <= char <= '\u4dbf')
            words = re.findall(r'[a-zA-Z]+', text)
            english_words_count = len(words)
            numbers = re.findall(r'[0-9]+', text)
            numbers_count = len(numbers)
            punctuation_count = sum(1 for char in text if unicodedata.category(char).startswith('P'))
            return chinese_count + english_words_count + numbers_count + punctuation_count
        else:  # "all"
            return sum(1 for char in text if not char.isspace())

    def _is_section_heading(self, text, section_patterns):
        if not text: return False
        for pattern in section_patterns:
            if re.fullmatch(pattern, text.strip(), re.IGNORECASE):
                return True
        return False

    def extract_sections(self):
        extracted_data = {}
        potential_headings = []
         # Prioritize main sections for matching
        main_section_keys = ["权利要求书", "说明书摘要", "说明书"]
        other_section_keys = [key for key in COMMON_SECTIONS_PATTERNS.keys() if key not in main_section_keys]
        ordered_section_keys = main_section_keys + other_section_keys


        for i, para_text in enumerate(self.paragraphs):
            for sec_key in ordered_section_keys: # Use the defined order for matching
                patterns = COMMON_SECTIONS_PATTERNS.get(sec_key, [])
                if self._is_section_heading(para_text, patterns):
                    # Avoid re-adding if a more general pattern (like '摘要' for '说明书摘要') matched first
                    # This needs careful ordering of COMMON_SECTIONS_PATTERNS or more complex logic
                    # For now, assume first match is good enough for a given paragraph
                    potential_headings.append({"name": sec_key, "index": i, "text": para_text})
                    break # A paragraph is only one heading type
        potential_headings.sort(key=lambda x: x["index"]) # Sort by appearance in doc

        if not potential_headings and self.paragraphs: # No standard headings found
            logger.warning("未能识别任何标准章节标题。将整个文档视为'全文内容'。")
            extracted_data["全文内容"] = {
                "content_paragraphs": self.paragraphs, "start_index": 0, "title_text": "全文内容 (未识别明确章节)"
            }
            return extracted_data

        for i in range(len(potential_headings)):
            current_heading = potential_headings[i]
            section_name = current_heading["name"]
            title_paragraph_text = current_heading["text"] # The text of the heading itself

            # Content starts AFTER the heading, unless it's a section where the heading IS the content (e.g. an abstract title)
            content_start_index = current_heading["index"] + 1
            # For sections like "摘要" or "说明书摘要", the heading itself might be the only line OR content starts immediately
            if section_name in ["说明书摘要", "摘要"]: # These sections' content usually starts on the same line or immediately after.
                 # If the title is the only thing, content_paragraphs will be empty from slicing, which is fine.
                 # If content follows title, it's covered.
                 # The current logic for content_paragraphs means it will take from next line.
                 # We might need to adjust if the title line ITSELF should be part of the content_paragraphs for these.
                 # For now, let's assume title is separate and content follows.
                 pass


            if i + 1 < len(potential_headings):
                next_heading_start_index = potential_headings[i+1]["index"]
                section_paras = self.paragraphs[content_start_index:next_heading_start_index]
            else: # Last heading
                section_paras = self.paragraphs[content_start_index:]

            # If a section is re-identified (e.g. '摘要' after '说明书摘要'), we might overwrite.
            # COMMON_SECTIONS_PATTERNS order and specificity is key.
            if section_name in extracted_data:
                logger.warning(f"章节 '{section_name}' (标题: '{title_paragraph_text}') 被多次识别或其模式与之前匹配的章节冲突。后匹配的将覆盖。")
            extracted_data[section_name] = {
                "content_paragraphs": section_paras,
                "start_index": current_heading["index"],
                "title_text": title_paragraph_text
            }
        return extracted_data

    def analyze(self):
        logger.info(f"开始分析文档 '{self.file_path.name}'")
        analysis_result = {
            "文件名": self.file_path.name,
            "统计模式": self.count_mode,
            "总字数": self.count_chars(self.full_text_content),
            "各部分": {},
            "检查结果": []
        }
        extracted_sections_data = self.extract_sections()
        actual_counts = {"总字数": analysis_result["总字数"]}

        for section_name, data in extracted_sections_data.items():
            content_text = "\n".join(data["content_paragraphs"])
            char_count = self.count_chars(content_text)
            analysis_result["各部分"][section_name] = {
                "字数": char_count,
                "原始内容标题": data.get("title_text", section_name),
            }
            actual_counts[section_name] = char_count
            logger.debug(f"章节 '{section_name}' (标题: '{data.get('title_text', '')}'): 字数 {char_count}")


        # Aggregation for "说明书"
        if "说明书" in self.config and "sub_sections" in self.config["说明书"]:
            sub_sections_total_chars = 0
            found_sub_sections = []
            logger.debug(f"说明书聚合检查: 配置子章节 {self.config['说明书']['sub_sections']}")
            for sub_sec_name in self.config["说明书"]["sub_sections"]:
                if sub_sec_name in actual_counts:
                    sub_sections_total_chars += actual_counts[sub_sec_name]
                    found_sub_sections.append(sub_sec_name)
                    logger.debug(f"聚合: 添加子章节 '{sub_sec_name}' 字数 {actual_counts[sub_sec_name]}")
                else:
                    logger.debug(f"聚合: 未找到子章节 '{sub_sec_name}' 的字数")


            if found_sub_sections: # Only update if sub-sections were actually found and summed
                actual_counts["说明书"] = sub_sections_total_chars # For checking requirements
                if "说明书" not in analysis_result["各部分"]: # If "说明书" itself wasn't a heading
                     analysis_result["各部分"]["说明书"] = {"字数": 0, "原始内容标题": "说明书 (聚合)", "包含子部分": []}

                analysis_result["各部分"]["说明书"]["字数"] = sub_sections_total_chars
                analysis_result["各部分"]["说明书"]["聚合字数"] = True
                analysis_result["各部分"]["说明书"]["包含子部分"] = found_sub_sections
                logger.info(f"说明书聚合计算完成: 总字数 {sub_sections_total_chars}, 包含 {found_sub_sections}")
            elif "说明书" not in actual_counts : # If "说明书" is in config but no subsections found AND no main "说明书" heading found
                logger.warning("配置了说明书聚合，但未找到任何定义的子章节，且未直接识别'说明书'章节。")


        # Check requirements
        logger.debug(f"开始检查要求，当前配置: {self.config}")
        for name, req in self.config.items():
            check_item = {"name": name, "actual": "N/A", "expected_str": "", "status_bool": None, "status_str": "", "message": ""}

            is_aggregated_section = name == "说明书" and analysis_result["各部分"].get(name, {}).get("聚合字数")

            if name not in actual_counts and "ratio" not in req and not is_aggregated_section :
                 # If it's "说明书" and it's supposed to be aggregated but wasn't, it means sub-sections weren't found.
                if name == "说明书" and "sub_sections" in req:
                    check_item["status_str"] = "未识别子章节"
                    check_item["message"] = f"{name}: 未能识别或聚合其子章节。"
                else:
                    check_item["status_str"] = "未识别"
                    check_item["message"] = f"{name}: 未在文档中识别到该部分。"
                check_item["expected_str"] = str(req)
                analysis_result["检查结果"].append(check_item)
                logger.debug(f"检查项 '{name}': {check_item['status_str']}")
                continue

            current_actual_chars = actual_counts.get(name, 0 if not is_aggregated_section else actual_counts.get(name,0) ) # Default to 0 if not found

            # Ratio check
            if "ratio" in req:
                target_section_name_for_ratio = name # e.g., "具体实施方式"
                ref_section_name_for_ratio = req["reference"] # e.g., "权利要求书"
                check_item["name"] = f"{target_section_name_for_ratio}/{ref_section_name_for_ratio} 比例"

                chars_target = actual_counts.get(target_section_name_for_ratio)
                chars_ref = actual_counts.get(ref_section_name_for_ratio)

                if chars_target is not None and chars_ref is not None and chars_ref > 0:
                    ratio_val = chars_target / chars_ref
                    target_ratio_val, tolerance_val = req["ratio"], req.get("tolerance", 0.1)
                    min_r_val, max_r_val = target_ratio_val * (1 - tolerance_val), target_ratio_val * (1 + tolerance_val)
                    check_item["status_bool"] = min_r_val <= ratio_val <= max_r_val
                    check_item["actual"] = f"{ratio_val:.2f}"
                    check_item["expected_str"] = f"{target_ratio_val:.2f} (±{tolerance_val*100:.0f}%, 即 {min_r_val:.2f}-{max_r_val:.2f})"
                    check_item["message"] = f"{check_item['name']}: {check_item['actual']} (目标: {check_item['expected_str']})"
                    # Check for additional min char count for the target section of ratio
                    if req.get("min") and chars_target < req.get("min"):
                        check_item["status_bool"] = False # Override status if min not met
                        check_item["message"] += f"；且 '{target_section_name_for_ratio}' 字数 {chars_target} 未达到最小要求 {req['min']}"
                elif chars_target is None:
                    check_item["status_str"] = "未识别目标章节"
                    check_item["message"] = f"进行比例计算时未识别到章节 '{target_section_name_for_ratio}'"
                else: # chars_ref is None or 0
                    check_item["status_str"] = "未识别参考章节或其字数为0"
                    check_item["message"] = f"进行比例计算时未识别到参考章节 '{ref_section_name_for_ratio}' 或其字数为0"
                analysis_result["检查结果"].append(check_item)
                logger.debug(f"检查项 '{check_item['name']}': {check_item['status_bool'] if check_item['status_bool'] is not None else check_item['status_str']}")
                continue


            # Absolute count check
            check_item["actual"] = current_actual_chars
            min_val_req, max_val_req = req.get("min"), req.get("max")
            message_suffix = ""

            if min_val_req is not None and max_val_req is not None:
                check_item["status_bool"] = min_val_req <= current_actual_chars <= max_val_req
                check_item["expected_str"] = f"{min_val_req}-{max_val_req}字"
            elif min_val_req is not None:
                check_item["status_bool"] = current_actual_chars >= min_val_req
                check_item["expected_str"] = f"至少{min_val_req}字"
            elif max_val_req is not None:
                check_item["status_bool"] = current_actual_chars <= max_val_req
                check_item["expected_str"] = f"不超过{max_val_req}字"
            else: # No specific min/max requirement for this section in config
                check_item["status_str"] = "信息" # Mark as informational
                check_item["expected_str"] = "无特定范围要求"
                # No status_bool needed as it's just info

            check_item["message"] = f"{name}: {current_actual_chars}字 (要求: {check_item['expected_str']})"

            if check_item["status_bool"] == False: # Check for explicit False
                if min_val_req is not None and current_actual_chars < min_val_req:
                    message_suffix = f" (字数不足，差 {min_val_req - current_actual_chars} 字)"
                elif max_val_req is not None and current_actual_chars > max_val_req:
                    message_suffix = f" (字数过多，多 {current_actual_chars - max_val_req} 字)"
            check_item["message"] += message_suffix
            analysis_result["检查结果"].append(check_item)
            logger.debug(f"检查项 '{name}': {check_item['status_bool'] if check_item['status_bool'] is not None else check_item['status_str']}")

        logger.info(f"分析完成 for '{self.file_path.name}'")
        return analysis_result

def get_default_config_yaml_str():
    return yaml.dump(DEFAULT_REQUIREMENTS, default_flow_style=False, allow_unicode=True, sort_keys=False)